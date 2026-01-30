import os
import sys
import re
import json
import statistics
import pandas as pd
import pdfplumber
from pypdf import PdfReader
import docx
from docx import Document
# PDF 生成库
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

# ==========================================
# 1. 全局配置与清洗模块 (Config & Cleaning)
# ==========================================

# 通用噪音正则 (去除了特定厂商名，保留通用特征)
COMMON_NOISE = [
    r'^\d+$',                      # 纯数字页码
    r'^Page \s*\d+',               # Page 10
    r'^\d+-\d+$',                  # 10-20
    r'^\.{4,}',                    # 目录省略号 ......
    r'^Copyright', r'^版权所有',    # 版权
    r'^文档版本', r'^Release',
]

def is_noise(text):
    """通用去噪"""
    text = text.strip()
    if not text: return True
    for pattern in COMMON_NOISE:
        if re.search(pattern, text, re.IGNORECASE):
            return True
    return False

def fix_broken_text(text):
    """
    语义修复 (完美版核心功能)
    1. 修复连字符: "con- \n figuration" -> "configuration"
    2. 修复伪换行: 只有以标点结尾的行才保留换行，否则合并
    """
    # 1. 修复连字符 (Hyphenation)
    # 匹配: 单词 + 连字符 + 可能的空格 + 换行 + 可能的空格 + 单词
    text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)

    # 2. 简单的行合并 (慎用，仅针对明显的句子断裂)
    # 逻辑：如果一行结尾不是 [. ! ? : ; 。！？：；] 且下一行开头是小写字母或中文，尝试合并
    # 这里为了稳健，我们先只做简单的去多余空格
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if not line: continue
        cleaned_lines.append(line)
    
    return "\n".join(cleaned_lines)

def table_to_markdown(table_data):
    """将 list-of-list 表格数据转换为 Markdown 格式字符串"""
    if not table_data: return ""
    
    # 过滤空行
    rows = [row for row in table_data if any(cell and cell.strip() for cell in row)]
    if not rows: return ""

    md_lines = []
    # 表头
    headers = [str(cell).replace('\n', ' ') if cell else "" for cell in rows[0]]
    md_lines.append("| " + " | ".join(headers) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    
    # 内容
    for row in rows[1:]:
        clean_row = [str(cell).replace('\n', '<br>') if cell else "" for cell in row]
        md_lines.append("| " + " | ".join(clean_row) + " |")
    
    return "\n" + "\n".join(md_lines) + "\n"

def extract_page_text_with_tables(pdf_page):
    """
    混合提取文本和表格 (优化版)
    先提取表格，将其转为 Markdown，再提取文本，尽量保持顺序
    """
    # 1. 提取表格
    tables = pdf_page.extract_tables()
    table_texts = [table_to_markdown(t) for t in tables]
    
    # 2. 提取纯文本
    raw_text = pdf_page.extract_text() or ""
    
    # 3. 简单的策略：将表格附加在文本后面 (或者尝试替换)
    # 由于 PDF 文本流和表格流很难完美穿插，最稳健的方式是：
    # 清洗文本中的表格残留碎片(难做)，或者直接把表格附在当页文本最后。
    # 这里采用：清洗后的文本 + 表格 Markdown
    
    cleaned_text = clean_text_block(raw_text)
    
    final_content = cleaned_text
    if table_texts:
        final_content += "\n\n【检测到参数表/数据表】:\n" + "\n".join(table_texts)
        
    return final_content

def clean_text_block(text):
    """基础清洗 + 语义修复"""
    text = fix_broken_text(text)
    lines = text.split('\n')
    cleaned = [line.strip() for line in lines if not is_noise(line)]
    return "\n".join(cleaned)

# ==========================================
# 2. 侦察与路由 (Scout & Router)
# ==========================================

def scout_file(file_path):
    """基于文档特征的智能路由"""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx': return "DOCX_STYLE"
    
    if ext == '.pdf':
        print(f"   🕵️ 正在分析文档结构特征: {os.path.basename(file_path)}")
        try:
            reader = PdfReader(file_path)
            # 特征 1: 目录树 (Tree)
            # 阈值：书签数量 > 15 且 层级深度往往由书签结构决定
            if reader.outline and len(reader.outline) > 15:
                print(f"      -> 特征识别: 完备目录树结构 (Tree Structure)")
                return "PDF_TREE"
        except: pass

        # 特征 2: 正则序列 (Sequential Regex)
        try:
            with pdfplumber.open(file_path) as pdf:
                sample_pages = pdf.pages[:10]
                regex_hits = 0
                total_lines = 0
                # 匹配 X.Y 或 X.Y.Z 格式
                pattern = re.compile(r"^\s*\d+\.\d+")
                for p in sample_pages:
                    text = p.extract_text() or ""
                    lines = text.split('\n')
                    total_lines += len(lines)
                    for line in lines:
                        if pattern.match(line.strip()):
                            regex_hits += 1
                
                density = regex_hits / total_lines if total_lines > 0 else 0
                if density > 0.05: # 5% 的行是编号标题
                    print(f"      -> 特征识别: 序列化编号结构 (Sequential/Regex)")
                    return "PDF_REGEX"
        except: pass
            
        # 特征 3: 默认视觉 (Visual)
        print(f"      -> 特征识别: 无明显结构，使用视觉字号分析 (Visual Layout)")
        return "PDF_VISUAL"
            
    return "UNKNOWN"

# ==========================================
# 3. 解析引擎群 (Engines)
# ==========================================

def engine_docx_style(file_path):
    print("   🚀 启动 DOCX 样式解析引擎...")
    doc = docx.Document(file_path)
    chunks = []
    current_chunk = {"title": "文档说明", "content": ""}
    
    # 泛化的样式匹配
    HEADER_KEYWORDS = ['heading', 'title', 'chapter', 'sect', '标题']
    REGEX_HEADER = re.compile(r"^\d+(\.\d+)*\s+")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if is_noise(text): continue
        
        style_name = para.style.name.lower()
        
        # 判定标题
        is_header = any(k in style_name for k in HEADER_KEYWORDS)
        # 补漏：Normal样式但长得像标题
        if not is_header and 'normal' in style_name and len(text) < 50:
             if REGEX_HEADER.match(text): is_header = True
        
        if is_header:
            if current_chunk["content"].strip(): chunks.append(current_chunk)
            # 上下文注入
            content_with_ctx = f"【导航路径】: {text}\n\n{text}"
            current_chunk = {"title": text, "content": content_with_ctx}
        else:
            current_chunk["content"] += text + "\n"
            
    if current_chunk["content"].strip(): chunks.append(current_chunk)
    return chunks

def engine_pdf_tree(file_path):
    print("   🚀 启动 PDF 递归目录解析引擎 (含表格重构)...")
    reader = PdfReader(file_path)
    chunks = []
    bookmarks = []
    
    # 稳健的递归提取
    def _flatten(outline, path=[]):
        idx = 0
        while idx < len(outline):
            item = outline[idx]
            if not isinstance(item, list):
                title = item.title
                try:
                    page = reader.get_destination_page_number(item) + 1
                    full_path = " > ".join(path + [title])
                    bookmarks.append({"title": full_path, "page": page, "raw_title": title})
                    # 检查子节点
                    if idx + 1 < len(outline) and isinstance(outline[idx+1], list):
                        _flatten(outline[idx+1], path + [title])
                        idx += 1
                except: pass
            idx += 1
    if reader.outline: _flatten(reader.outline)

    print(f"      提取到 {len(bookmarks)} 个节点，开始深度解析...")
    
    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)
        for i, node in enumerate(bookmarks):
            start = node['page']
            end = bookmarks[i+1]['page'] if i < len(bookmarks)-1 else total_pages
            
            # 限制跨度，防止内存溢出
            if end - start > 20: end = start + 20 
            
            content_buffer = ""
            for p_idx in range(start-1, end): 
                 if p_idx < len(pdf.pages):
                     # 使用混合提取 (文本+表格)
                     page_content = extract_page_text_with_tables(pdf.pages[p_idx])
                     content_buffer += page_content + "\n"
            
            if content_buffer.strip():
                # 上下文注入 (Context Injection)
                full_content = f"【导航路径】: {node['title']}\n\n{content_buffer.strip()}"
                chunks.append({"title": node['title'], "content": full_content})
                
    return chunks

def engine_pdf_regex(file_path):
    print("   🚀 启动 PDF 正则序列解析引擎...")
    PATTERN = re.compile(r"^\s*(\d+(\.\d+)+)\s+(.*)")
    chunks = []
    current_chunk = {"title": "前言", "content": ""}
    
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # 简单文本提取 (对于纯列表文档，表格通常不是重点，且容易打断正则)
            text = clean_text_block(page.extract_text() or "")
            lines = text.split('\n')
            for line in lines:
                match = PATTERN.match(line)
                if match:
                    if current_chunk["content"].strip(): chunks.append(current_chunk)
                    
                    raw_title = f"{match.group(1)} {match.group(3)}"
                    # 上下文注入
                    full_content = f"【指令节点】: {raw_title}\n\n{line}"
                    current_chunk = {"title": raw_title, "content": full_content}
                else:
                    current_chunk["content"] += line + "\n"
                    
    if current_chunk["content"].strip(): chunks.append(current_chunk)
    return chunks

def engine_pdf_visual(file_path):
    print("   🚀 启动 PDF 视觉布局解析引擎...")
    chunks = []
    current_chunk = {"title": "Start", "content": ""}
    
    with pdfplumber.open(file_path) as pdf:
        # 计算基准
        sample_sizes = []
        for p in pdf.pages[:5]:
            words = p.extract_words(extra_attrs=["size"])
            sample_sizes.extend([w['size'] for w in words])
        base_size = statistics.median(sample_sizes) if sample_sizes else 10
        threshold = base_size + 1.5 
        
        for page in pdf.pages:
            # 简单按行聚合
            words = page.extract_words(extra_attrs=["size", "top"])
            lines_dict = {}
            for w in words:
                top = round(w['top'])
                if top not in lines_dict: lines_dict[top] = []
                lines_dict[top].append(w)
            
            sorted_tops = sorted(lines_dict.keys())
            for top in sorted_tops:
                line_words = sorted(lines_dict[top], key=lambda x: x['x0'])
                text = " ".join([w['text'] for w in line_words])
                max_size = max([w['size'] for w in line_words])
                
                if is_noise(text): continue
                
                if max_size >= threshold and len(text) < 80:
                    if current_chunk["content"].strip(): chunks.append(current_chunk)
                    full_content = f"【章节】: {text}\n\n{text}"
                    current_chunk = {"title": text, "content": full_content}
                else:
                    current_chunk["content"] += text + "\n"
                    
    if current_chunk["content"].strip(): chunks.append(current_chunk)
    return chunks

# ==========================================
# 4. 多格式输出器 (Multi-Format Writers)
# ==========================================

def save_as_excel(data, path):
    df = pd.DataFrame(data)
    # 过滤掉内部字段
    cols = [c for c in ['title', 'content', 'filename', 'parsing_method'] if c in df.columns]
    df[cols].to_excel(path, index=False)
    print(f"      [Excel] 已生成: {path}")

def save_as_json(data, path):
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"      [JSON]  已生成: {path}")

def save_as_word(data, path):
    doc = Document()
    doc.add_heading("文档提取结果", 0)
    
    for item in data:
        # 标题
        doc.add_heading(item.get('title', 'Untitled'), level=1)
        # 元数据
        meta = f"Source: {item.get('filename')} | Method: {item.get('parsing_method')}"
        p = doc.add_paragraph(meta)
        p.style = "Intense Quote"
        # 内容
        doc.add_paragraph(item.get('content', ''))
        doc.add_paragraph("-" * 20)
        
    doc.save(path)
    print(f"      [Word]  已生成: {path}")

def save_as_pdf(data, path):
    if not HAS_REPORTLAB:
        print("      [PDF] ❌ 无法生成: 缺少 'reportlab' 库")
        return

    # 尝试查找中文字体
    font_path = None
    possible_paths = [
        "C:/Windows/Fonts/simhei.ttf",  # Windows 黑体
        "C:/Windows/Fonts/msyh.ttf",    # Windows 微软雅黑
        "/System/Library/Fonts/PingFang.ttc", # Mac
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf" # Linux
    ]
    for p in possible_paths:
        if os.path.exists(p):
            font_path = p
            break
            
    if not font_path:
        print("      [PDF] ⚠️ 跳过生成: 未找到中文字体 (simhei/msyh)。")
        return

    try:
        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
        c = canvas.Canvas(path, pagesize=A4)
        width, height = A4
        y_pos = height - 50
        
        c.setFont("ChineseFont", 10)
        
        for item in data:
            # 简单的 PDF 写入逻辑 (不处理复杂的自动换行，仅作演示)
            title = f"Title: {item.get('title', '')}"
            content = item.get('content', '')[:200] + "..." # 只写前200字避免溢出
            
            if y_pos < 100: 
                c.showPage()
                y_pos = height - 50
                c.setFont("ChineseFont", 10)
            
            c.drawString(50, y_pos, title)
            y_pos -= 20
            
            # 简单的多行处理
            lines = content.split('\n')
            for line in lines:
                if y_pos < 50:
                    c.showPage()
                    y_pos = height - 50
                    c.setFont("ChineseFont", 10)
                c.drawString(50, y_pos, line)
                y_pos -= 15
            y_pos -= 20
            
        c.save()
        print(f"      [PDF]   已生成: {path}")
    except Exception as e:
        print(f"      [PDF]   生成出错: {e}")

# ==========================================
# 5. 主流程 (Main)
# ==========================================

def process_file_full_suite(file_path, output_dir, export_formats, base_filename=None):
    filename = os.path.basename(file_path)
    if not base_filename:
        base_filename = os.path.splitext(filename)[0]
    
    print(f"\n[{filename}] 开始处理...")
    
    # 1. 路由
    engine_type = scout_file(file_path)
    
    # 2. 执行
    data = []
    try:
        if engine_type == "DOCX_STYLE": data = engine_docx_style(file_path)
        elif engine_type == "PDF_TREE": data = engine_pdf_tree(file_path)
        elif engine_type == "PDF_REGEX": data = engine_pdf_regex(file_path)
        elif engine_type == "PDF_VISUAL": data = engine_pdf_visual(file_path)
        else:
            print("   ⚠️ 无法识别文件类型。")
            return
    except Exception as e:
        print(f"   ❌ 解析出错: {e}")
        return

    if not data:
        print("   ⚠️ 未提取到数据。")
        return
        
    # 注入元数据
    for item in data:
        item['filename'] = filename
        item['parsing_method'] = engine_type

    # 3. 多格式导出
    print(f"   💾 正在导出数据 ({len(data)} 条)...")
    base_path = os.path.join(output_dir, base_filename)
    
    if 'xlsx' in export_formats: save_as_excel(data, base_path + ".xlsx")
    if 'json' in export_formats: save_as_json(data, base_path + ".json")
    if 'docx' in export_formats: save_as_word(data, base_path + ".docx")
    if 'pdf' in export_formats:  save_as_pdf(data, base_path + ".pdf")

def main():
    print("="*60)
    print("   🌌 通用智能文档解析器 V3.0 (Ultimate Edition)")
    print("   特性: 表格重构 | 语义修复 | 多格式输出 | 去厂商化")
    print("="*60)
    
    # 1. 输入文件
    files = []
    if len(sys.argv) > 1:
        path = sys.argv[1]
    else:
        raw = input("\n👉 [1/3] 请拖入文件或文件夹: ").strip()
        path = raw.replace('"', '').replace("'", "")
    
    if os.path.isfile(path): files.append(path)
    elif os.path.isdir(path):
        for f in os.listdir(path):
            if f.lower().endswith(('.pdf', '.docx')):
                files.append(os.path.join(path, f))
                
    if not files: print("❌ 无有效文件"); return

    # 2. 输出目录
    default_out = "parsed_results"
    raw_out = input(f"\n👉 [2/3] 输出目录 (默认 '{default_out}'): ").strip()
    out_dir = raw_out.replace('"', '').replace("'", "") if raw_out else default_out
    if not os.path.exists(out_dir): os.makedirs(out_dir)

    # 3. 格式选择
    print("\n👉 [3/3] 选择输出格式 (多选用逗号分隔, 默认 Excel+JSON)")
    print("   可用: xlsx, json, docx, pdf")
    fmt_in = input("   格式: ").strip().lower()
    
    formats = []
    if not fmt_in: 
        formats = ['xlsx', 'json'] # 默认
    else:
        if 'xlsx' in fmt_in: formats.append('xlsx')
        if 'json' in fmt_in: formats.append('json')
        if 'docx' in fmt_in: formats.append('docx')
        if 'pdf' in fmt_in:  formats.append('pdf')
        
    # 执行
    print("\n🚀 任务队列启动...")
    for f in files:
        # 允许用户自定义单文件名前缀? 这里简化为自动命名
        process_file_full_suite(f, out_dir, formats)
        
    print(f"\n🎉 全部完成! 查看目录: {out_dir}")

if __name__ == "__main__":
    main()